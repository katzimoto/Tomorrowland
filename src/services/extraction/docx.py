"""DOCX text extractor."""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from services.extraction.base import ExtractionResult, LocationSegment


def _is_heading(paragraph: object) -> bool:
    """Return True when *paragraph* uses a heading style."""
    try:
        # python-docx paragraphs have .style.name — accept any duck-typed object
        name = (paragraph.style.name or "") if paragraph.style else ""  # type: ignore[attr-defined]
        return name.startswith("Heading")
    except (AttributeError, ValueError):
        return False


class DocxExtractor:
    """Extract text from Word .docx files using python-docx."""

    def extract(self, path: Path) -> ExtractionResult:
        """Return concatenated text from all paragraphs and tables.

        python-docx yields merged cells once per spanned column, so a cell
        that spans 3 columns appears 3 times in ``row.cells``.  We deduplicate
        by the underlying ``_tc`` XML element identity to produce each cell's
        text exactly once.

        Location segments capture *section_heading* from heading-style paragraphs
        (style.name starts with ``"Heading"``).
        """
        try:
            doc = Document(str(path))
            texts: list[str] = []
            segments: list[LocationSegment] = []
            current_heading: str | None = None
            segment_start: int = 0

            # --- First pass: paragraphs ---
            for p in doc.paragraphs:
                if not p.text:
                    continue
                offset = len("\n".join(texts)) if texts else 0
                if offset > 0:
                    offset += 1  # newline separator

                if _is_heading(p):
                    # Close previous heading's content segment
                    if current_heading is not None and segment_start < offset:
                        segments.append(
                            LocationSegment(
                                start_char=segment_start,
                                end_char=offset,
                                section_heading=current_heading,
                            )
                        )
                    # Record heading text as its own segment
                    current_heading = p.text
                    segment_start = offset
                    heading_end = offset + len(p.text)
                    segments.append(
                        LocationSegment(
                            start_char=offset,
                            end_char=heading_end,
                            section_heading=current_heading,
                        )
                    )
                    segment_start = heading_end + 1

                texts.append(p.text)

            # Close trailing content under the last heading
            if current_heading is not None:
                full_len = len("\n".join(texts)) if texts else 0
                if segment_start < full_len:
                    segments.append(
                        LocationSegment(
                            start_char=segment_start,
                            end_char=full_len,
                            section_heading=current_heading,
                        )
                    )

            # --- Second pass: tables (appended after paragraphs) ---
            for table in doc.tables:
                for row in table.rows:
                    seen_tc: set[int] = set()
                    for cell in row.cells:
                        tc_id = id(cell._tc)  # noqa: SLF001
                        if tc_id in seen_tc:
                            continue
                        seen_tc.add(tc_id)
                        if cell.text:
                            texts.append(cell.text)

            return ExtractionResult(
                text="\n".join(texts),
                location_segments=segments,
            )
        except (OSError, KeyError, ValueError, zipfile.BadZipFile, PackageNotFoundError):
            return ExtractionResult(text="")
